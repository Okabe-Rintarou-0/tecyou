import argparse
import datetime
import os
import sys

from PyQt6.QtWidgets import QApplication, QFileDialog
from PIL import Image
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def read_wh(image_file):
    image = Image.open(image_file)
    width, height = image.size
    image_width = Cm(3)
    image_height = Cm(4)
    if width > height:
        image_width, image_height = image_height, image_width
    image.close()
    return image_width, image_height


def create_word_document(images_dir, output_dir):
    # 创建 Word 文档
    document = Document()

    # 获取所有图片文件
    image_files = get_image_files(images_dir)

    # 计算每个图片的宽度和高度
    i = 0
    tot = len(image_files)
    max_cols = 4
    while i < tot:
        image_file = image_files[i]
        width, height = read_wh(image_file)
        row_len = width.cm
        # 插入图片到表格单元格中
        table = document.add_table(rows=1, cols=max_cols)
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        run = cell.paragraphs[0].add_run()
        run.add_picture(image_file, width=width, height=height)
        added = 1
        for j in range(1, max_cols):
            if i + j < tot:
                cell = table.cell(0, j)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                run = cell.paragraphs[0].add_run()
                width, height = read_wh(image_files[i + j])
                row_len += width.cm
                added += 1
                run.add_picture(image_files[i + j], width=width, height=height)
            if row_len >= 15:
                break

        i += added

    # 保存 Word 文档
    output_file = f"output_{datetime.datetime.now().strftime('%Y%m%d')}.docx"
    output_path = os.path.join(output_dir, output_file)
    document.save(output_path)
    print(f"Word 文档已保存到：{output_path}")


def get_image_files(images_dir):
    image_files = []
    for file_name in os.listdir(images_dir):
        file_path = os.path.join(images_dir, file_name)
        if os.path.isfile(file_path):
            # 仅筛选常见图片格式的文件
            if file_name.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp")):
                image_files.append(file_path)
    return image_files


if __name__ == "__main__":
    # 初始化Qt应用程序
    app = QApplication(sys.argv)

    # 使用QFileDialog获取目标目录
    target_directory = QFileDialog.getExistingDirectory(
        None, "选择目标目录", ".", QFileDialog.ShowDirsOnly
    )

    # 如果用户取消选择目标目录，退出程序
    if not target_directory:
        sys.exit()

    # 使用QFileDialog获取输出目录
    output_directory = QFileDialog.getExistingDirectory(
        None, "选择输出目录", ".", QFileDialog.ShowDirsOnly
    )

    # 如果用户取消选择输出目录，退出程序
    if not output_directory:
        sys.exit()

    # 调用你的处理函数
    create_word_document(target_directory, output_directory)

    # 程序执行完毕，退出应用程序
    sys.exit(app.exec_())
